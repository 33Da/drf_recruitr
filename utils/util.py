from rest_framework.views import exception_handler
from rest_framework.exceptions import APIException
from rest_framework import status
from django.http import HttpResponse


def custom_exception_handler(exc, context):
    response = exception_handler(exc, context)  # 获取本来应该返回的exception的response

    if response is not None:
        response.data['status_code'] = response.status_code  # 可添加status_code
        # response.data['error_code'] = 1
        try:
            response.data["message"] = response.data['detail']  # 增加message这个key
            del response.data['detail']
        except:
            pass
    if response is None:
        return HttpResponse("禁止单独测试接口")

    return response


class myException422(APIException):
    status_code = status.HTTP_422_UNPROCESSABLE_ENTITY


def jwt_response_username_userid_token(token, user=None, request=None):
    """
    自定义验证成功后的返回数据处理函数
    :param token:
    :param user:
    :param request:
    :return:
    """

    if user.check == 0:
        data = {
            # jwt令牌
            'status_code': 200,
            "message":"等待管理员审核账号"
        }
    else:
        data = {
            # jwt令牌
            'status_code': 200,
            'token': token,
            'user_id': user.id,
            'username': user.username,
            'role': user.role
        }

    return data





import xlwt,datetime
from xlwt import *
import xlrd

# 读取excle
def read_to_excle(file):
    """读取excle表中第二列准考证号"""
    myWorkbook = xlrd.open_workbook(file)

    mySheets = myWorkbook.sheets()  # 获取工作表list。

    mySheet = mySheets[0]  # 通过索引顺序获取。

    # 获取行数
    nrows = mySheet.nrows

    examcodes = []
    for row in range(1,nrows):
        myCellValue = mySheet.cell_value(row, 1)
        examcodes.append(myCellValue)

    return examcodes





# 写入excel文件函数
def wite_to_excel(name,head_data,records):
    #获取时间戳
    timestr = datetime.datetime.now().strftime("%Y%m%d%H%M%S")

    # 工作表
    wbk = xlwt.Workbook()
    sheet1 = wbk.add_sheet('sheet1',cell_overwrite_ok=True)

    #写入表头
    for filed in range(0,len(head_data)):
        sheet1.write(0,filed,head_data[filed],excel_head_style())


    #写入数据记录
    for row in range(1,len(records)+1):

        for col in range(0,len(head_data)):
            sheet1.write(row,col,records[row-1][col],excel_record_style())
            #设置默认单元格宽度
            sheet1.col(col).width = 256*15

    cur_path = os.path.abspath('.')
    # 设置生成文件所在路径
    download_url = cur_path + '/media/data/download/'

    wbk.save(download_url+name+'-'+timestr+'.xls')
    return '/media/data/download/' + name +'-'+timestr+'.xls'

# 定义导出文件表头格式
def excel_head_style():
    # 创建一个样式
    style = XFStyle()
    #设置背景色
    pattern = Pattern()
    pattern.pattern = Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = Style.colour_map['light_green']  # 设置单元格背景色
    style.pattern = pattern


    # 设置字体
    # font0 = xlwt.Font()
    # font0.name = u'微软雅黑'
    # font0.bold = True
    # font0.colour_index = 0
    # font0.height = 240
    # style.font = font0
    #设置文字位置
    alignment = xlwt.Alignment()  # 设置字体在单元格的位置
    alignment.horz = xlwt.Alignment.HORZ_CENTER  # 水平方向
    alignment.vert = xlwt.Alignment.VERT_CENTER  # 竖直方向
    style.alignment = alignment
    # 设置边框
    borders = xlwt.Borders()  # Create borders
    borders.left = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.right = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.top = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.bottom = xlwt.Borders.THIN  # 添加边框-虚线边框
    style.borders = borders



    return style

# 定义导出文件记录格式
def excel_record_style():
    # 创建一个样式
    style = XFStyle()
    #设置字体
    font0 = xlwt.Font()
    font0.name = u'微软雅黑'
    font0.bold = False
    font0.colour_index = 0
    font0.height = 200
    style.font = font0
    #设置文字位置
    alignment = xlwt.Alignment()  # 设置字体在单元格的位置
    alignment.horz = xlwt.Alignment.HORZ_CENTER  # 水平方向
    alignment.vert = xlwt.Alignment.VERT_CENTER  # 竖直方向
    style.alignment = alignment
    # 设置边框
    borders = xlwt.Borders()  # Create borders
    borders.left = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.right = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.top = xlwt.Borders.THIN  # 添加边框-虚线边框
    borders.bottom = xlwt.Borders.THIN  # 添加边框-虚线边框
    style.borders = borders

    return style



# 获取招生人
import zipfile   # 改了源码
import os
import docx
from docx import Document
from drf_recruitr.settings import MEDIA_ROOT

def un_zip(file_name,dir):
    """unzip zip file"""

    try:
        zip_file = zipfile.ZipFile(file_name)
        for names in zip_file.namelist():

            zip_file.extract(names, dir)
        zip_file.close()
    except Exception as e:
        print(e)
        return False
    return True



def read_ordinaryword():
    """读取本科基本信息"""

    dir = MEDIA_ROOT + "/upload/ordinary/本科考生电子档案/"
    result = []
    try:
        students = os.listdir(dir)
    except Exception as e:
        print(e)
        return result

    try:

        for student in students:
            # 读取基本信息
            stu_dir = dir + student + "/基本信息.docx"
            document1 = Document(stu_dir)  # 读入文件
            tables = document1.tables  # 获取表格
            table = tables[0]

            student_dict = {}

            student_dict["name"] = table.cell(0, 0).text.split("：")[1]
            student_dict["idcard"] = table.cell(4, 0).text.split("：")[1]
            student_dict["address"] = table.cell(6, 0).text.split("：")[1]
            student_dict["polit_physcs"] = table.cell(14, 0).text.split("：")[1]
            student_dict["history_chemistry"] = table.cell(15, 0).text.split("：")[1]
            student_dict["geography_biology"] = table.cell(15, 1).text.split("：")[1]


            subject = table.cell(7, 0).text.split("：")[1]

            if subject == "文科类":
                student_dict["subject"] = 0
            else:
                student_dict["subject"] = 1




            if "江门" in student_dict["address"]:
                student_dict["is_city"] = True
            else:
                student_dict["is_city"] = False

            if "广东" in student_dict["address"]:
                student_dict["is_province"] = True
            else:
                student_dict["is_province"] = False

            # 读取志愿信息
            stu_dir = dir + student + "/志愿信息.docx"
            document2 = Document(stu_dir)  # 读入文件
            tables = document2.tables  # 获取表格
            table = tables[0]

            student_dict["first_expectation"] = table.cell(2, 6).text
            student_dict["second_expectation"] = table.cell(2, 7).text
            student_dict["third_exception"] = table.cell(2, 8).text
            student_dict["fourth_expectation"] = table.cell(2, 9).text
            student_dict["fifth_expectation"] = table.cell(2, 10).text
            student_dict["sixth_expectation"] = table.cell(2, 11).text
            if table.cell(2, 12).text == "是":
                student_dict["is_dispensing"] = True
            else:
                student_dict["is_dispensing"] = False


            # 读取科目信息
            stu_dir = dir + student + "/高考成绩.docx"
            document3 = Document(stu_dir)  # 读入文件
            tables = document3.tables  # 获取表格
            table = tables[0]
            student_dict["chinese"] = int(table.cell(5, 0).text.split("：")[1])
            student_dict["math"] = int(table.cell(5, 1).text.split("：")[1])
            student_dict["english"] = int(table.cell(6, 0).text.split("：")[1])
            student_dict["integrated_subject"] = int(table.cell(6, 1).text.split("：")[1])
            student_dict["total"] = int(table.cell(4, 0).text.split("：")[1])
            student_dict["examcode"] = table.cell(3, 0).text.split("：")[1]

            student_dict["is_pass"] = False
            student_dict["expectation"] = 0

            result.append(student_dict)

    except Exception as e:
       print(e)
    finally:
        return result


def read_postgraduateword():
    """读取研究生基本信息"""

    dir = MEDIA_ROOT + "/upload/postgraduate/firstexam/考生电子档案/"
    result = []
    try:
        students = os.listdir(dir)
    except Exception as e:
        print(e)
        return result

    try:

        for student in students:
            # 读取基本信息
            stu_dir = dir + student
            document = Document(stu_dir)  # 读入文件
            tables = document.tables  # 获取表格
            table1 = tables[0]

            student_dict = {}

            student_dict["name"] = table1.cell(0, 0).text.split("：")[1]
            student_dict["idcard"] = table1.cell(4, 0).text.split("：")[1]
            student_dict["examcode"] = table1.cell(5, 0).text.split("：")[1]
            student_dict["phone"] = table1.cell(7, 0).text.split("：")[1]

            table2 = tables[1]
            student_dict["politics"] = int(table2.cell(1, 0).text.split("：")[1])
            student_dict["base_subject"] = int(table2.cell(2, 0).text.split("：")[1])
            student_dict["english"] = int(table2.cell(1, 1).text.split("：")[1])
            student_dict["speciality_subject"] = int(table2.cell(2, 1).text.split("：")[1])
            student_dict["total"] = int(table2.cell(0, 0).text.split("：")[1])

            table3 = tables[2]
            student_dict["want_expectation"] = int(table3.cell(1, 1).text)
            student_dict["research_interests"] = table3.cell(1, 3).text
            student_dict["teacher"] = table3.cell(1, 4).text

            student_dict["is_pass"] = False
            student_dict["expectation"] = None
            student_dict["passfirstexam"] = False
            student_dict["passsecondexam"] = False

            result.append(student_dict)

    except Exception as e:
       print(e)
    finally:
        return result


def ordinary_recruit():

    if un_zip(MEDIA_ROOT + "/upload/ordinary/考生电子档案.zip",MEDIA_ROOT + "/upload/ordinary/"):
        data = read_ordinaryword()
    else:
        data = {}
    return data


def postgraduate_recruit():

    if un_zip(MEDIA_ROOT + "/upload/postgraduate/firstexam/考生电子档案.zip",MEDIA_ROOT + "/upload/postgraduate/firstexam/"):
        data = read_postgraduateword()
    else:
        data = {}
    return data

def read_specialword():
    """专插本逻辑"""
    dir = MEDIA_ROOT + "/upload/SpecialStudent/exam/考生电子档案/"
    result = []
    try:
        students = os.listdir(dir)
    except Exception as e:
        print(e)
        return result

    try:

        for student in students:
            # 读取基本信息
            stu_dir = dir + student
            document = Document(stu_dir)  # 读入文件
            tables = document.tables  # 获取表格
            table1 = tables[0]

            student_dict = {}

            student_dict["name"] = table1.cell(0, 0).text.split("：")[1]
            student_dict["idcard"] = table1.cell(4, 0).text.split("：")[1]
            student_dict["phone"] = table1.cell(7, 0).text.split("：")[1]
            student_dict["address"] = table1.cell(6, 0).text.split("：")[1]


            student_dict["is_pass"] = False
            student_dict["file"] = "/media/upload/SpecialStudent/考生电子档案/" + student

            result.append(student_dict)

    except Exception as e:
        print(e)
    finally:
        return result

def read_specialexamword():
    """读取专插本考试成绩"""

    dir = MEDIA_ROOT + "/upload/SpecialStudent/exam/考生电子档案/"
    result = []
    try:
        students = os.listdir(dir)
    except Exception as e:
        print(e)
        return result
    try:
        for student in students:
            # 读取基本信息
            stu_dir = dir + student

            document = Document(stu_dir)  # 读入文件
            tables = document.tables  # 获取表格

            table1 = tables[0]

            student_dict = {}

            student_dict["name"] = table1.cell(0, 0).text.split("：")[1]
            student_dict["idcard"] = table1.cell(2, 0).text.split("：")[1]
            student_dict["examcode"] = table1.cell(3, 0).text.split("：")[1]

            table2 = tables[1]
            student_dict["total"] = table2.cell(0, 0).text.split("：")[1]
            student_dict["polity"] = table2.cell(1, 0).text.split("：")[1]
            student_dict["english"] = table2.cell(1, 1).text.split("：")[1]
            student_dict["speciality_subject1"] = table2.cell(3, 0).text.split("：")[1]
            student_dict["speciality_subject2"] = table2.cell(3, 1).text.split("：")[1]
            student_dict["speciality_basesubject"] = table1.cell(2, 0).text.split("：")[1]


            table3 = tables[2]
            student_dict["want_expectation"] = table3.cell(1, 1).text


            student_dict["is_pass"] = False


            result.append(student_dict)

    except Exception as e:
        print(e)
    finally:
        return result

def special_recriit():
    if un_zip(MEDIA_ROOT + "/upload/SpecialStudent/考生电子档案.zip", MEDIA_ROOT + "/upload/SpecialStudent/"):
        data = read_specialword()
    else:
        data = {}
    return data

def special_examrecriit():
    if un_zip(MEDIA_ROOT + "/upload/SpecialStudent/exam/考生电子档案.zip", MEDIA_ROOT + "/upload/SpecialStudent/exam/"):
        data = read_specialexamword()
    else:
        data = {}
    return data